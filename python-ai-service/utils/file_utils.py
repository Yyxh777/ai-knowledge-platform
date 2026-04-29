import os
import fitz  # PyMuPDF：pip install pymupdf
import requests
from docx import Document
from io import BytesIO


def download_file_from_url(url: str) -> bytes:
    """从 URL 下载文件，返回二进制内容。"""
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        return response.content
    except requests.exceptions.RequestException as e:
        raise Exception(f"下载文件失败: {str(e)}")


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    用 PyMuPDF（fitz）从 PDF 提取文本。

    相比 PyPDF2 的优势：
    - 对中文 PDF 的字符提取准确率更高（底层使用 MuPDF 引擎）
    - 可正确处理多栏排版、嵌入字体、扫描页等复杂场景
    - 长期活跃维护，PyPDF2 已停止更新（最后版本 3.0.1，2023 年归档）
    """
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        pages = []
        for page in doc:
            text = page.get_text("text")  # "text" 模式：按阅读顺序提取纯文本
            if text.strip():
                pages.append(text)
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        raise Exception(f"解析PDF文件失败: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """从 .docx 文件中提取文本（包含正文段落和表格内容）。"""
    try:
        doc = Document(BytesIO(file_content))
        parts = []

        # 正文段落
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # 表格内容（原版本未处理，此处补充）
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)
    except Exception as e:
        raise Exception(f"解析DOCX文件失败: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> str:
    """从文本文件提取内容，按优先级尝试多种编码。"""
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for encoding in encodings:
        try:
            return file_content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise Exception("无法识别文件编码")


def get_file_extension(url: str) -> str:
    """从 URL 中提取文件扩展名（小写，不含点号）。"""
    url_without_params = url.split("?")[0]
    ext = os.path.splitext(url_without_params)[1].lower()
    return ext.lstrip(".")
